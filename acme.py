# importing packages
from docxtpl import DocxTemplate, InlineImage
import docxtpl
import pandas as pd
import numpy as np 
import re
from docx.shared import Mm
import docx
from docx.enum.text import WD_BREAK
from pandas.core.indexes.base import Index
import logging 

# log
logging.basicConfig(filename="Aurabot.log", format='%(asctime)s %(message)s', filemode='w')
logger = logging.getLogger()
logger.setLevel(logging.DEBUG)
logger.info("Started")

try:
    iserror = []
    df = pd.read_excel(r'C:\Users\ezgi turalı\Desktop\occasional works\aurabot\Acme\ReportExcel.xlsx', sheet_name='Header')

    # Nan to empty
    df['Comment English'].replace(np.nan, '', inplace=True)
    df['Comment Turkish'].replace(np.nan, '', inplace=True)

    language = len(df['Comment Turkish'][0])

    if language > 0:
        desc = pd.read_excel(r'C:\Users\ezgi turalı\Desktop\occasional works\aurabot\Acme\ReportExcel.xlsx', sheet_name='Description')
        doc = DocxTemplate(r'C:\Users\ezgi turalı\Desktop\occasional works\aurabot\Acme\Acme_Rapor_Format_TR.docx')
    else:
        desc = pd.read_excel(r'C:\Users\ezgi turalı\Desktop\occasional works\aurabot\Acme\ReportExcel.xlsx', sheet_name='Description')
        df = pd.read_excel(r'C:\Users\ezgi turalı\Desktop\occasional works\aurabot\Acme\ReportExcel.xlsx', sheet_name='Header')
        doc = DocxTemplate(r'C:\Users\ezgi turalı\Desktop\occasional works\aurabot\Acme\Acme_Rapor_Format_EN.docx')

    # title column has whitespacing at the end
    desc = desc.rename(columns={'Title ': 'Title'})

    for text_idx, text in enumerate(desc['Description']):
        text = text.strip()
        text = text.replace('>', '·       ')
        text = text.replace('\n>', '·       ')
        desc['Description'][text_idx] = text

    # splitting 'Description' column and making new variables
    if  language > 0:
        new = desc['Description'].str.split("Gözlem:", expand=True)
        new.columns = ['A', 'Gozlem']

        new2 = new["Gozlem"].str.split("Risk:", expand = True)
        new2.columns = ['Gozlem', 'Risk']

        new3 = new2["Risk"].str.split("Öneri:", expand = True)
        new3.columns = ['Risk', 'Öneri']

        descr = new2['Gozlem']
        descr = pd.merge(descr, new3, left_index=True, right_index=True)
    else:
        new = desc['Description'].str.split("Observation:", expand=True)
        new.columns = ['A', 'Observation']

        new2 = new["Observation"].str.split("Risk:", expand = True)
        new2.columns = ['Observation', 'Risk']

        new3 = new2["Risk"].str.split("Recommendation:", expand = True)
        new3.columns = ['Risk', 'Recommendation']

        descr = new2['Observation']
        descr = pd.merge(descr, new3, left_index=True, right_index=True)

    # merging with main dataset
    desc.drop('Description', axis=1, inplace=True)
    desc = pd.merge(desc, descr, how="outer", left_index=True, right_index=True)

    # excluding yönetim görüşü so that we can merge to the main dataset
    if  language > 0:
        com = df['Comment Turkish']
        com = com.str.split("Yönetim Görüşü:", expand=True)
        com.columns=['Görüş', 'YönetimGörüşü']
        com.drop('Görüş', axis=1, inplace=True)
        
    else: 
        com = df['Comment English']
        com = com.str.split("Management Comment:", expand=True)
        com.columns=['Observation', 'ManagementComment']
        com.drop('Observation', axis=1, inplace=True)

    # merging all columns into one dataframe
    desc = pd.merge(desc, com, how="outer", left_index=True, right_index=True)

    # repating n times yönetim görüşü
    if  language > 0:
        desc["YönetimGörüşü"].replace(np.nan, desc["YönetimGörüşü"][0], inplace=True)
    else: 
        desc["ManagementComment"].replace(np.nan, desc["ManagementComment"][0], inplace=True)

    # merging image 
    img1 = [InlineImage(doc, 'C:\\Users\\TR101909\\Desktop\\aurabot\\high_risk.png', width = Mm(20)), 
            InlineImage(doc, 'C:\\Users\\TR101909\\Desktop\\aurabot\\low_risk.png', width = Mm(40))]
    img = pd.DataFrame(img1)
    img.columns=['img']

    desc = pd.merge(desc, img, how="outer", left_index=True, right_index=True)
        
    # if condition to classification
    desc['image'] = np.where(desc['Classification'] == 'Significant', desc['img'][0], desc['img'][1])

    # title numbering
    desc['number'] = desc['Risk'].index + 1
        
    # dropping old img column
    desc.drop('img', axis= 1, inplace= True)

    # defining variables
    header = df.iloc[0]['Header']
    desc_dict = desc.to_dict('records')
            
    # making dictionary for render
    context = {
        'header':header,
        'desc': desc_dict,
            }

    # rendering document
    doc.render(context)

    # saving document
    doc.save(str(header) + """ .docx""")
        

    # adjusting sections with docx
    doc2 = docx.Document(str(header) + """ .docx""")

    for paragraph in doc2.paragraphs:
        if '-----' in paragraph.text:
            run = paragraph.add_run()
            run.add_break(WD_BREAK.PAGE)

    doc2.save(str(header) + """ .docx""")
    print("Rapor Üretildi.")
    logger.info("No problem encountered")
    error_file = open(r'C:\Users\ezgi turalı\Desktop\occasional works\aurabot\Acme\error_log.txt','w')
    error_file.writelines("No problem encountered")
    
except Exception as ex:
    iserror.append(str(ex))
    error_file = open(r'C:\Users\ezgi turalı\Desktop\occasional works\aurabot\Acme\error_log.txt','w')
    error_file.writelines(iserror)
    logger.info(ex)